# -*- coding: utf-8 -*-
import io
import re
import html
import camelot
import pdfplumber
import requests
import streamlit as st
from PIL import Image
from docx import Document
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import landscape
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    SimpleDocTemplate, LongTable, TableStyle,
    Paragraph, Spacer, Image as RLImage
)
import openai

# ─── OpenAI Setup ─────────────────────────────────────────────────────────────
openai.api_key = st.secrets["OPENAI_API_KEY"]
GPT_MODEL = "gpt-4o-mini"

# ─── Font Registration ────────────────────────────────────────────────────────
try:
    pdfmetrics.registerFont(TTFont("DMSerif", "fonts/DMSerifDisplay-Regular.ttf"))
    pdfmetrics.registerFont(TTFont("Barlow", "fonts/Barlow-Regular.ttf"))
    DEFAULT_SERIF_FONT = "DMSerif"
    DEFAULT_SANS_FONT  = "Barlow"
except:
    DEFAULT_SERIF_FONT = "Times New Roman"
    DEFAULT_SANS_FONT  = "Arial"

# ─── Streamlit Boilerplate ────────────────────────────────────────────────────
st.set_page_config(page_title="Proposal Transformer", layout="wide")
st.title("🔄 Proposal Layout Transformer")
st.write("Upload a vertically formatted proposal PDF and download both PDF and Word outputs.")
uploaded = st.file_uploader("Upload proposal PDF", type="pdf")
if not uploaded:
    st.stop()
pdf_bytes = uploaded.read()

# ─── Helpers ──────────────────────────────────────────────────────────────────
def split_cell_text(raw: str):
    lines = [l.strip() for l in raw.splitlines() if l.strip()]
    if not lines:
        return "", ""
    strat = lines[0]
    desc  = " ".join(lines[1:])
    desc  = re.sub(r'\s+', ' ', desc).strip()
    return strat, desc

def add_hyperlink(paragraph, url, text, font_name=None, font_size=None, bold=None):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    styles = part.document.styles
    if "Hyperlink" not in styles:
        style = styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER, True)
        style.font.color.rgb     = RGBColor(0x05, 0x63, 0xC1)
        style.font.underline      = True
        style.priority            = 9
        style.unhide_when_used    = True
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    if font_name:
        rf = OxmlElement("w:rFonts")
        rf.set(qn("w:ascii"), font_name)
        rf.set(qn("w:hAnsi"), font_name)
        rPr.append(rf)
    if font_size:
        sz   = OxmlElement("w:sz");   sz.set(qn("w:val"), str(int(font_size * 2)))
        szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), str(int(font_size * 2)))
        rPr.append(sz); rPr.append(szCs)
    if bold:
        b = OxmlElement("w:b"); rPr.append(b)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return docx.text.run.Run(new_run, paragraph)

def ai_extract_header(page_image_bytes):
    resp = openai.chat.completions.create(
        model=GPT_MODEL,
        inputs=[{"image": page_image_bytes}],
        query="""
Please OCR the top table's header row on this PDF page and return exactly
the eight column names in JSON array format, in order.
""",
        response_format="json"
    )
    return resp.value

# ─── Extract Tables ────────────────────────────────────────────────────────────
tables_info = []
grand_total = None
proposal_title = "Untitled Proposal"

with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
    texts = [p.extract_text() or "" for p in pdf.pages]
    first_lines = texts[0].splitlines()
    title_cand = next((l for l in first_lines if "proposal" in l.lower()), None)
    proposal_title = title_cand.strip() if title_cand else first_lines[0].strip()

    page1 = pdf.pages[0]
    img = page1.to_image(resolution=150).original
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    try:
        ai_hdr = ai_extract_header(buf.read())
    except:
        ai_hdr = None

    used_totals = set()
    def find_total(pi):
        if pi >= len(texts): return None
        for ln in texts[pi].splitlines():
            if re.search(r'\b(?!grand\s)total\b.*?\$\s*[\d,.]+', ln, re.I) and ln not in used_totals:
                used_totals.add(ln)
                return ln.strip()
        return None

    for pi, page in enumerate(pdf.pages):
        if pi == 0:
            try:
                cams = camelot.read_pdf(io.BytesIO(pdf_bytes), pages="1", flavor="lattice", strip_text="\n")
                df   = cams[0].df
                data = df.values.tolist()
                source = "camelot"
            except:
                tbls    = page.find_tables()
                tbl_obj = tbls[0]
                data    = tbl_obj.extract(x_tolerance=1, y_tolerance=1)
                source  = "plumber"
        else:
            tbls = page.find_tables()
            if not tbls:
                continue
            tbl_obj = tbls[0]
            data    = tbl_obj.extract(x_tolerance=1, y_tolerance=1)
            source  = "plumber"

        if not data or len(data) < 2:
            continue

        raw_hdr = [str(h).strip() for h in data[0]]
        if pi == 0 and ai_hdr and len(ai_hdr) == 8:
            hdr = ai_hdr
        else:
            hdr = [h for h in raw_hdr if h and h.lower() != "none"]
            if len(hdr) != 8:
                hdr = (hdr[:8] + [""]*8)[:8]
        hdr[0], hdr[1] = "Strategy", "Description"
        desc_i = hdr.index("Description")

        rows_data = []
        row_links = []
        table_total = None

        if source == "plumber":
            desc_links = {}
            links = page.hyperlinks or []
            for rid, row_obj in enumerate(tbl_obj.rows):
                if rid == 0: continue
                if desc_i < len(row_obj.cells):
                    bbox = row_obj.cells[desc_i]
                    if not bbox: continue
                    x0, top, x1, bottom = bbox
                    for lk in links:
                        if all(k in lk for k in ("x0","x1","top","bottom","uri")):
                            if not (lk["x1"] < x0 or lk["x0"] > x1 or lk["bottom"] < top or lk["top"] > bottom):
                                desc_links[rid] = lk["uri"]
                                break

            for rid, row in enumerate(data[1:], start=1):
                cells = [str(c).strip() for c in row]
                if not any(cells): continue
                lower0 = cells[0].lower()
                if ("total" in lower0 or "subtotal" in lower0) and any("$" in c for c in cells):
                    if table_total is None:
                        table_total = cells
                    continue

                strat, desc = split_cell_text(cells[desc_i] if desc_i < len(cells) else "")
                rest = cells[:desc_i] + cells[desc_i+1:]
                words = desc.split()
                chunk_size = 20
                desc_chunks = [ ' '.join(words[i:i+chunk_size]) for i in range(0, len(words), chunk_size) ] or [""]

                for ci, chunk in enumerate(desc_chunks):
                    if ci == 0:
                        rows_data.append([strat, chunk] + rest)
                        row_links.append(desc_links.get(rid))
                    else:
                        rows_data.append(["", chunk] + [""] * len(rest))
                        row_links.append(None)
        else:
            for row in data[1:]:
                cells = [str(c).strip() for c in row]
                if not any(cells): continue
                lower0 = cells[0].lower()
                if ("total" in lower0 or "subtotal" in lower0) and any("$" in c for c in cells):
                    if table_total is None:
                        table_total = cells
                    continue

                strat, desc = split_cell_text(cells[1] if len(cells) > 1 else "")
                rest = cells[2:]
                words = desc.split()
                chunk_size = 20
                desc_chunks = [ ' '.join(words[i:i+chunk_size]) for i in range(0, len(words), chunk_size) ] or [""]

                for ci, chunk in enumerate(desc_chunks):
                    if ci == 0:
                        rows_data.append([strat, chunk] + rest)
                        row_links.append(None)
                    else:
                        rows_data.append(["", chunk] + [""] * len(rest))
                        row_links.append(None)

        if table_total is None:
            table_total = find_total(pi)

        desired = ["Strategy","Description","Start Date","End Date","Term (Months)","Monthly Amount","Item Total","Notes"]
        idx_map = {h: i for i, h in enumerate(hdr)}

        new_rows = []
        for r in rows_data:
            new_rows.append([ r[idx_map[col]] if col in idx_map and idx_map[col] < len(r) else "" for col in desired ])

        if isinstance(table_total, list):
            new_tot = [ table_total[idx_map[col]] if col in idx_map and idx_map[col] < len(table_total) else "" for col in desired ]
        else:
            new_tot = table_total

        tables_info.append((desired, new_rows, row_links, new_tot))

    for tx in reversed(texts):
        m = re.search(r'Grand\s+Total.*?(\$\s*[\d,]+\.\d{2})', tx, re.I|re.S)
        if m:
            grand_total = m.group(1).replace(" ", "")
            break

# ─── BUILD PDF ─────────────────────────────────────────────────────────────────
pdf_buf = io.BytesIO()
doc = SimpleDocTemplate(
    pdf_buf,
    pagesize=landscape((17*inch,11*inch)),
    leftMargin=0.5*inch,
    rightMargin=0.5*inch,
    topMargin=0.5*inch,
    bottomMargin=0.5*inch,
)
ts = ParagraphStyle("Title",  fontName=DEFAULT_SERIF_FONT, fontSize=18, alignment=TA_CENTER, spaceAfter=12)
hs = ParagraphStyle("Header", fontName=DEFAULT_SERIF_FONT, fontSize=10, alignment=TA_CENTER, textColor=colors.black)
bs = ParagraphStyle("Body",   fontName=DEFAULT_SANS_FONT,  fontSize=9,  alignment=TA_LEFT,   leading=11)
bl = ParagraphStyle("BL",     fontName=DEFAULT_SERIF_FONT, fontSize=10, alignment=TA_LEFT,   spaceBefore=6)
br = ParagraphStyle("BR",     fontName=DEFAULT_SERIF_FONT, fontSize=10, alignment=TA_RIGHT,  spaceBefore=6)

elements = []
try:
    logo_url = "https://www.carnegiehighered.com/wp-content/uploads/2021/11/Twitter-Image-2-2021.png"
    resp     = requests.get(logo_url, timeout=10); resp.raise_for_status()
    logo     = resp.content
    img      = Image.open(io.BytesIO(logo))
    ratio    = img.height / img.width
    w        = min(5*inch, doc.width)
    h        = w * ratio
    elements.append(RLImage(io.BytesIO(logo), width=w, height=h))
except:
    pass

elements += [Spacer(1,12), Paragraph(html.escape(proposal_title), ts), Spacer(1,24)]
total_w = doc.width

for hdr, rows, links, tot in tables_info:
    n = len(hdr)
    desc_idx = hdr.index("Description")
    desc_w   = total_w * 0.45
    other_w  = (total_w - desc_w) / (n - 1)
    col_ws   = [desc_w if i == desc_idx else other_w for i in range(n)]

    wrapped = [[Paragraph(html.escape(h), hs) for h in hdr]]
    for i, row in enumerate(rows):
        line = []
        for j, cell in enumerate(row):
            txt = html.escape(cell)
            if j == desc_idx and i < len(links) and links[i]:
                line.append(Paragraph(f"{txt} <link href='{html.escape(links[i])}' color='blue'>- link</link>", bs))
            else:
                line.append(Paragraph(txt, bs))
        wrapped.append(line)

    if tot:
        lbl, val = "Total", ""
        if isinstance(tot, list):
            lbl = tot[0] or "Total"
            val = next((c for c in reversed(tot) if "$" in c), "")
        else:
            m = re.match(r'(.*?)\s*(\$[\d,]+\.\d{2})', tot)
            if m:
                lbl, val = m.group(1).strip(), m.group(2)
        wrapped.append([Paragraph(lbl, bl)] + [Paragraph("", bs)]*(n-2) + [Paragraph(val, br)])

    tbl = LongTable(wrapped, colWidths=col_ws, repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0), colors.HexColor("#F2F2F2")),
        ("GRID",(0,0),(-1,-1),0.25,colors.grey),
        ("VALIGN",(0,0),(-1,0),"MIDDLE"),
        ("VALIGN",(0,1),(-1,-1),"TOP"),
    ]))
    elements += [tbl, Spacer(1,24)]

if grand_total and tables_info:
    hdr = tables_info[-1][0]
    n   = len(hdr)
    desc_idx = hdr.index("Description")
    desc_w   = total_w * 0.45
    other_w  = (total_w - desc_w) / (n - 1)
    col_ws   = [desc_w if i == desc_idx else other_w for i in range(n)]

    row = [Paragraph("Grand Total", bl)] + [Paragraph("", bs)]*(n-2) + [Paragraph(html.escape(grand_total), br)]
    gt  = LongTable([row], colWidths=col_ws)
    gt.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),colors.HexColor("#E0E0E0")),
        ("GRID",(0,0),(-1,-1),0.25,colors.grey),
        ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ("SPAN",(0,0),(-2,0)),
        ("ALIGN",(0,0),(-2,0),"LEFT"),
        ("ALIGN",(-1,0),(-1,0),"RIGHT"),
    ]))
    elements += [gt]

doc.build(elements)
pdf_buf.seek(0)

# ─── BUILD WORD ───────────────────────────────────────────────────────────────
docx_buf = io.BytesIO()
docx_doc = Document()
sec = docx_doc.sections[0]
sec.orientation    = WD_ORIENT.LANDSCAPE
sec.page_width     = Inches(17)
sec.page_height    = Inches(11)
sec.left_margin    = Inches(0.5)
sec.right_margin   = Inches(0.5)
sec.top_margin     = Inches(0.5)
sec.bottom_margin  = Inches(0.5)

if 'logo' in locals():
    try:
        p = docx_doc.add_paragraph(); r = p.add_run()
        img = Image.open(io.BytesIO(logo))
        w_in  = 5
        r.add_picture(io.BytesIO(logo), width=Inches(w_in))
        p.alignment = WD_TABLE_ALIGNMENT.CENTER
    except:
        pass

pt = docx_doc.add_paragraph(); pt.alignment = WD_TABLE_ALIGNMENT.CENTER
rt = pt.add_run(proposal_title); rt.font.name = DEFAULT_SERIF_FONT; rt.font.size = Pt(18); rt.bold = True
docx_doc.add_paragraph()

TOTAL_W_IN = sec.page_width.inches - sec.left_margin.inches - sec.right_margin.inches

for hdr, rows, links, tot in tables_info:
    n = len(hdr)
    if n < 1:
        continue
    desc_idx = hdr.index("Description")
    desc_w   = 0.45 * TOTAL_W_IN
    other_w  = (TOTAL_W_IN - desc_w) / (n - 1)

    tbl = docx_doc.add_table(rows=1, cols=n, style="Table Grid")
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.allow_autofit = False
    tbl.autofit       = False

    tblPr_list = tbl._element.xpath("./w:tblPr")
    if not tblPr_list:
        tblPr = OxmlElement("w:tblPr"); tbl._element.insert(0, tblPr)
    else:
        tblPr = tblPr_list[0]
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:w"), "5000"); tblW.set(qn("w:type"), "pct")
    existing = tblPr.xpath("./w:tblW")
    if existing:
        tblPr.remove(existing[0])
    tblPr.append(tblW)

    for i, col in enumerate(tbl.columns):
        col.width = Inches(desc_w if i == desc_idx else other_w)

    hdr_cells = tbl.rows[0].cells
    for i, name in enumerate(hdr):
        cell = hdr_cells[i]
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd"); shd.set(qn("w:fill"), "F2F2F2"); tcPr.append(shd)
        p = cell.paragraphs[0]; p.text = ""
        r = p.add_run(name); r.font.name = DEFAULT_SERIF_FONT; r.font.size = Pt(10); r.bold = True
        p.alignment = WD_TABLE_ALIGNMENT.CENTER; cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for ridx, row in enumerate(rows):
        rcells = tbl.add_row().cells
        for cidx, val in enumerate(row):
            cell = rcells[cidx]
            p = cell.paragraphs[0]; p.text = ""
            run = p.add_run(str(val)); run.font.name = DEFAULT_SANS_FONT; run.font.size = Pt(9)
            if cidx == desc_idx and ridx < len(links) and links[ridx]:
                p.add_run(" "); add_hyperlink(p, links[ridx], "- link", font_name=DEFAULT_SANS_FONT, font_size=9)
            p.alignment = WD_TABLE_ALIGNMENT.LEFT; cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    if tot:
        trow = tbl.add_row().cells
        lbl, amt = "Total", ""
        if isinstance(tot, list):
            lbl = tot[0] or "Total"; amt = next((c for c in reversed(tot) if "$" in c), "")
        else:
            m = re.match(r'(.*?)\s*(\$[\d,]+\.\d{2})', tot)
            if m:
                lbl, amt = m.group(1).strip(), m.group(2)
        lc = trow[0]
        if n > 1:
            lc.merge(trow[n-2])
        p = lc.paragraphs[0]; p.text = ""; r = p.add_run(lbl)
        r.font.name = DEFAULT_SERIF_FONT; r.font.size = Pt(10); r.bold = True
        p.alignment = WD_TABLE_ALIGNMENT.LEFT; lc.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        ac = trow[n-1]; p2 = ac.paragraphs[0]; p2.text = ""
        r2 = p2.add_run(amt); r2.font.name = DEFAULT_SERIF_FONT; r2.font.size = Pt(10); r2.bold = True
        p2.alignment = WD_TABLE_ALIGNMENT.RIGHT; ac.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    docx_doc.add_paragraph()

if grand_total and tables_info:
    hdr = tables_info[-1][0]; n = len(hdr)
    desc_idx = hdr.index("Description")
    desc_w   = 0.45 * TOTAL_W_IN
    other_w  = (TOTAL_W_IN - desc_w) / (n - 1)

    tblg = docx_doc.add_table(rows=1, cols=n, style="Table Grid")
    tblg.alignment = WD_TABLE_ALIGNMENT.CENTER; tblg.allow_autofit = False; tblg.autofit = False
    tblPr_list = tblg._element.xpath("./w:tblPr")
    if not tblPr_list:
        tblPr = OxmlElement("w:tblPr"); tblg._element.insert(0, tblPr)
    else:
        tblPr = tblPr_list[0]
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:w"), "5000"); tblW.set(qn("w:type"), "pct")
    existing = tblPr.xpath("./w:tblW")
    if existing:
        tblPr.remove(existing[0])
    tblPr.append(tblW)

    for i, col in enumerate(tblg.columns):
        col.width = Inches(desc_w if i == desc_idx else other_w)

    cells = tblg.rows[0].cells
    lc = cells[0]
    if n > 1:
        lc.merge(cells[n-2])
    tc = lc._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "E0E0E0"); tcPr.append(shd)
    p = lc.paragraphs[0]; p.text = ""; r = p.add_run("Grand Total")
    r.font.name = DEFAULT_SERIF_FONT; r.font.size = Pt(10); r.bold = True
    p.alignment = WD_TABLE_ALIGNMENT.LEFT; lc.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    ac = cells[n-1]; tc2 = ac._tc; tcPr2 = tc2.get_or_add_tcPr()
    shd2 = OxmlElement("w:shd"); shd2.set(qn("w:fill"), "E0E0E0"); tcPr2.append(shd2)
    p2 = ac.paragraphs[0]; p2.text = ""
    r2 = p2.add_run(grand_total); r2.font.name = DEFAULT_SERIF_FONT; r2.font.size = Pt(10); r2.bold = True
    p2.alignment = WD_TABLE_ALIGNMENT.RIGHT; ac.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    docx_doc.add_paragraph()

docx_buf = io.BytesIO()
docx_doc.save(docx_buf)
docx_buf.seek(0)

c1, c2 = st.columns(2)
c1.download_button(
    "📥 Download deliverable PDF",
    data=pdf_buf,
    file_name="proposal_deliverable.pdf",
    mime="application/pdf",
    use_container_width=True
)
c2.download_button(
    "📥 Download deliverable DOCX",
    data=docx_buf,
    file_name="proposal_deliverable.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    use_container_width=True
)
